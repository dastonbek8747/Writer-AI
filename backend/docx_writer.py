import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def clean_filename(topic: str) -> str:
    name = re.sub(r'[\\/*?:"<>|]', "", topic)
    name = name.strip()
    return name[:120] if len(name) > 120 else name


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def setup_document_styles(document: Document):
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(14)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_cover_page(document: Document, topic: str):
    for _ in range(6):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KURS ISHI / MUSTAQIL ISH")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "Times New Roman"

    document.add_paragraph()

    mavzu = document.add_paragraph()
    mavzu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mavzu.add_run(f"MAVZU: {topic}")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "Times New Roman"

    for _ in range(12):
        document.add_paragraph()

    footer_p = document.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Synora AI")
    run.font.size = Pt(12)
    run.font.italic = True

    document.add_page_break()


def add_contents_page(document: Document, tasks: list):
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("MUNDARIJA")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = "Times New Roman"

    document.add_paragraph()

    for i, task in enumerate(tasks):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i + 1}. {task}")
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

    document.add_page_break()


def add_footer_page_number(document: Document):
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p)


def add_body_content(document: Document, tasks: list, results: list):
    for i in range(len(tasks)):
        heading = document.add_heading(f"{i + 1}. {tasks[i]}", level=1)

        paragraphs = [p.strip() for p in results[i].split("\n") if p.strip()]
        for paragraph_text in paragraphs:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            run = p.add_run(paragraph_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

        if i != len(tasks) - 1:
            document.add_page_break()


def generate_docx(topic: str, tasks: list, results: list) -> str:
    document = Document()
    setup_document_styles(document)

    add_cover_page(document, topic)
    add_contents_page(document, tasks)
    add_body_content(document, tasks, results)
    add_footer_page_number(document)

    os.makedirs("./Files", exist_ok=True)
    file_name = clean_filename(topic) + ".docx"
    file_path = os.path.join("./Files", file_name)
    document.save(file_path)

    return file_path
